import pandas as pd
import requests
import json
from pandas import ExcelWriter
import sys,os
from docx import Document
import pathlib
import re

class KBAutomatedTester:
    def __init__(self):
        self._url = "http://<ip address>/searchkb"
        self._payload = """{\"text\" : \"@@@@\",\"client\": \"xxxx\",\"user_name\": \"hari.radha\",
                            \"is_sfb\" : \"N\",\"is_initial_chat\" : \"N\"}"""
        self._headers = {'content-type': "application/json"}
        self._file_path = "Payload.xlsx"
        self._upload_url = "http://10.87.2.39/storeinkb"
        self._upload_payload = {
            "Subject": "","Body": "","Sender": "Null","DocumentType": "Generic","Recipients": "","CarbonCopy": "","BackCarbonCopy": "",
            "SentOn": "","ContributedBy": "thameem.sakkarai","Attachments": "","FileName": "@@@@","Title": "#### ",
            "ExistingDocObjectId": "","Client": "thameem","DocumentTypeExtension": "$$$$","CustomTags": [],"StaticQuestions": [] }
        self._upload_filepath="./InputDocs/"

    def upload_doc(self):
        return_message={}
        try:
            fullText = []
            for i in os.listdir(self._upload_filepath):
                if pathlib.Path(self._upload_filepath + "/" + str(i)).suffix == ".docx":
                    doc = Document(self._upload_filepath + i)
                    filename, extension = os.path.splitext(i)
                    extension = extension.upper()[1:]
                    for para in doc.paragraphs:
                        # if len(i.text) > 0:
                        #     line = re.sub('\s', ' ', i.text)
                        #     line = re.sub('“|”|‘|’', '"', line)
                        #     line = re.sub(' – ', '-', line)
                        #     fullText.append(re.findall('\w.+', line)[0])
                    # for para in doc.paragraphs:
                    #     fullText.append(re.sub('\‘|\’|\“|\”', '"', para.text))
                        fullText.append(para.text)
                elif pathlib.Path(self._upload_filepath + "/" + str(i)).suffix == ".txt":
                    filename, extension = os.path.splitext(i)
                    file = open(self._upload_filepath + i, "r")
                    fulltext = file.readlines()
                else:
                    return_message["message"]="Please provide docx/txt files"
                if(len(fullText) > 0):
                    print('try....')
                    input_payload = self._upload_payload
                    input_payload["Body"] = ' '.join(fullText)
                    # input_payload = input_payload.replace("&&&&", ' '.join(fullText))
                    # input_payload = input_payload.replace("$$$$", str(extension.upper()))
                    # input_payload = input_payload.replace("####", str(filename)).replace("@@@@", str(i))


                    print("----")
                    input_payload = json.dumps(input_payload)
                    print(input_payload)
                    response = requests.request("POST", url=self._upload_url, data=input_payload, headers=self._headers)
                    print(response)
                    if response.status_code == 200:
                        resp = json.dumps(response.text)
                        print(resp)
                    else:
                        if(response is not None):
                            return_message["message"] = "Server returned status: " + response.status_code
                        else:
                            return_message["message"] = "Could not get any response from server."
                else:
                    return_message["message"] = "No content found."
        except Exception as e:
            return_message["message"] = str(e)
        return return_message

    def get_test_payload(self):
        return_message = {}
        try:
            lst_payload_data = []
            df_data = pd.read_excel(self._file_path)
            for i in df_data.index:
                payload_item = dict()
                payload_item["Query"] = df_data["Query"][i]
                payload_item["ExpectedSummary"] = df_data["ExpectedSummary"][i]
                payload_item["ExpectedDetails"] = df_data["ExpectedDetails"][i]
                lst_payload_data.append(payload_item)
        except Exception as e:
            return_message["message"]=str(e)
        return lst_payload_data,return_message

    def search_kb_with_payload(self, lst_query_details):
        return_message={}
        try:
            for payload_item in lst_query_details:
                input_payload = self._payload
                input_payload = input_payload.replace("@@@@", str(payload_item["Query"].encode("utf-8", errors="strict")))
                response = requests.request("POST", url=self._url, data=input_payload, headers=self._headers)
                if response is not None and response.status_code == 200:
                    resp = json.loads(response.text)
                    payload_item["ActualSummary"] = resp["summary"]
                    payload_item["ActualDetails"] = resp["details"]
                    payload_item["StandardResponse"] = resp["standard_response"]
                    if payload_item["StandardResponse"] == "Y":
                        payload_item["Response"] = resp["response"]
        except Exception as e:
            return_message["message"] = str(e)
        return lst_query_details, return_message

    def log_kb_results(self, lst_data_details):
        df = pd.DataFrame(lst_data_details, columns=["Query", "ExpectedSummary", "ActualSummary", "ExpectedDetails",
                                                            "ActualDetails", "Passed", "StandardResponse", "Response",])
        writer = ExcelWriter("Output.xlsx", engine="xlsxwriter")
        df.to_excel(writer, sheet_name="Output",index=False)
        writer.save()

    def compare_response(self, lst_data):
        for data_item in lst_data:
            is_pass = False
            actual_summary = data_item["ActualSummary"]
            expected_summary = data_item["ExpectedSummary"]
            actual_details = data_item["ActualDetails"]
            expected_details = data_item["ExpectedDetails"]
            if actual_summary == expected_summary and actual_details == expected_details :
                is_pass = True
            if is_pass is True:
                data_item["Passed"] = "Yes"
            else:
                data_item["Passed"] = "No"
        return lst_data


if __name__ == "__main__":
    kb_automated_tester = KBAutomatedTester()
    kb_automated_tester.upload_doc()
    # lst_payload,return_message = kb_automated_tester.get_test_payload()
    # lst_data, return_message = kb_automated_tester.search_kb_with_payload(lst_payload)
    # lst_data = kb_automated_tester.compare_response(lst_data)
    # kb_automated_tester.log_kb_results(lst_data)